#!/usr/bin/python
# -*- coding: utf-8 -*-
import os
import sys
import importlib
import logging
import time
import datetime

import xlrd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert(config):

    # Open the excel file
    path_read = os.path.join(os.getcwd(), config['read_params']['source_name'])
    wb = xlrd.open_workbook(path_read, formatting_info=True)

    # Get the list of the sheets name
    sheet_list = wb.sheet_names()
    # Select one sheet and get its size
    s = wb.sheet_by_name(sheet_list[0])  # or s = wb.sheet_by_index(1)
    print("row and column number of excel: ", s.nrows, s.ncols)

    # Open the word template file
    path_template = os.path.join(os.getcwd(), config['read_params']['template_name'])
    doc_template = Document(path_template)
    #tables = doc_template.tables
    # for tabidx, table in enumerate(tables):
    #     print('table {}:\n'.format(tabidx), file=open("debug.txt", "a"))
    #     for rowidx, row in enumerate(table.rows):
    #         print('row {}:\n'.format(rowidx), file=open("debug.txt", "a"))
    #         for cellidx, cell in enumerate(row.cells):
    #             print('col {}:\n'.format(cellidx), file=open("debug.txt", "a"))
    #             for paragraph in cell.paragraphs:
    #                 print('{} | '.format(paragraph.text), file=open("debug.txt", "a", encoding="utf-8"))

    # Create sub_working_dir
    now = datetime.datetime.now()
    sub_working_dir = '{}/{}/{}'.format(
        os.getcwd(), config['output_dir'],
        time.strftime("%d%H%M%S", time.localtime()))
    if not os.path.exists(sub_working_dir):
        os.makedirs(sub_working_dir)
    logging.info("sub working dir: %s" % sub_working_dir)

    for i in range(s.nrows - 1):
        document = doc_template
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal'].font.size = Pt(10.5)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        table = document.tables[0]

        # 流水号
        table.cell(0, 14).text = str(s.cell(i+1, 1).value)
        table.cell(0, 14).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 送检单位
        table.cell(1, 1).text = str(s.cell(i+1, 0).value)
        table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(8)

        # 计量器具名称
        table.cell(1, 14).text = '压力表'
        table.cell(1, 14).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 制造单位
        table.cell(2, 1).text = str(s.cell(i+1, 2).value)
        table.cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 型号规格
        measure = str(s.cell(i+1, 3).value)
        index_suff = measure.lower().find('mpa')
        index_mid = measure.find('～') or measure.find('~')
        # output_measure = measure[:index_suff] + ' ' + measure[index_suff:]
        output_measure = measure[0] + ' ' + measure[1:index_mid] + ' ～ ' + measure[index_mid + 1:index_suff-1] + ' ' + measure[index_suff-1] + ' ' + measure[index_suff:]
        # measure.replace("", "  ")
        table.cell(3, 1).text = output_measure
        table.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 测量范围
        table.cell(4, 1).text = output_measure
        table.cell(4, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 最大允许误差
        table.cell(4, 14).text = '±     0.016 MPa      '
        table.cell(4, 14).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


        # 准确度等级
        # insert_paragraph_before(table.cell(5, 1), '1.0 ')
        tmp = table.cell(5, 1).text
        table.cell(5, 1).text = '         ' + str(s.cell(i+1, 5).value) + '   ' + tmp

        # 分度值
        table.cell(5, 14).text = '0.05 MPa      '
        table.cell(5, 14).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 出厂编号
        no_ = str(s.cell(i+1, 4).value)
        table.cell(6, 1).text = no_.rstrip('0').rstrip('.') if '.0' in no_ else no_         # incase change the number by program
        table.cell(6, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 环境温度
        table.cell(6, 8).text = str(s.cell(i+1, 6).value) + '   ℃'
        table.cell(6, 8).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 相对湿度
        tmp = table.cell(6, 17).text
        table.cell(6, 16).text = str(s.cell(i+1, 7).value) + '   %'
        table.cell(6, 16).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # if config['date_params'] == 0:
        parag = document.paragraphs[3]
        text = parag.text
        text_day = text.find('日')
        text_month = text.find('月')
        text_year = text.find('年')
        parag._p.clear()
        parag.add_run(text[:text_year-8] + ' ' + str(now.year) + '   ' + text[text_year] + '   ' + 
                    str(now.month) + '   ' + text[text_month] + '   ' + str(now.day) + '  ' + text[text_day])
        parag.runs[0].font.size = Pt(12)

        path_write = os.path.join(sub_working_dir, config['output_dir'] + str(i + 1) + '.docx')
        document.save(path_write)

def main():
    logging.basicConfig(level=logging.DEBUG,
                        format="[%(asctime)s %(filename)s] %(message)s")

    if len(sys.argv) != 2:
        logging.error("Usage: python training.py params.py")
        sys.exit()
    params_path = sys.argv[1]
    if not os.path.isfile(params_path):
        logging.error("no params file found! path: {}".format(params_path))
        sys.exit()
    config = importlib.import_module(params_path[:-3]).PARAMS
    
    convert(config)

if __name__ == "__main__":
    main()    