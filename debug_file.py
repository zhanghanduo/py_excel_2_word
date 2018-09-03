#!/usr/bin/python
# -*- coding: utf-8 -*-
import os
import sys
import importlib
import logging

from docx import Document

def debug_table(config):

    # Open the word template file
    path_template = os.path.join(os.getcwd(), config['read_params']['template_name'])
    doc_template = Document(path_template)
    tables = doc_template.tables
    for tabidx, table in enumerate(tables):
        print('table {}:\n'.format(tabidx), file=open("debug.txt", "a"))
        for rowidx, row in enumerate(table.rows):
            print('row {}:\n'.format(rowidx), file=open("debug.txt", "a"))
            for cellidx, cell in enumerate(row.cells):
                print('col {}:\n'.format(cellidx), file=open("debug.txt", "a"))
                for paragraph in cell.paragraphs:
                    print('{} | '.format(paragraph.text), file=open("debug.txt", "a", encoding="utf-8"))


def debug_others(config):

    # Open the word template file
    path_template = os.path.join(os.getcwd(), config['read_params']['template_name'])
    doc_template = Document(path_template)
    for index, parag in enumerate(doc_template.paragraphs):
        print('paragraph {}: '.format(index), file=open("debug1.txt", "a"))
        print('{} '.format(parag.text), file=open("debug1.txt", "a", encoding="utf-8"))

def main():
    logging.basicConfig(level=logging.DEBUG,
                        format="[%(asctime)s %(filename)s] %(message)s")

    if len(sys.argv) != 2:
        logging.error("Usage: python training.py params.py")
        sys.exit()
    params_path = sys.argv[1]
    if not os.path.isfile(params_path):
        logging.error("no params file found! path: {}".format(params_path))
        sys.exit()
    config = importlib.import_module(params_path[:-3]).PARAMS
    
    debug_others(config)

if __name__ == "__main__":
    main()    